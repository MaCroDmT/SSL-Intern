import tkinter as tk
from tkinter import messagebox, simpledialog
import threading
import pyautogui
import time
import os
from datetime import datetime, date

# --- Imports for Word Document ---
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: 'python-docx' library not found.")
    print("Please install it using: pip install python-docx")
    exit()

# --- Import for Outlook/Word Automation ---
try:
    import win32com.client
except ImportError:
    print("Error: 'pywin32' library not found.")
    print("Please install it using: pip install pywin32")
    exit()
# ----------------------------------------


class ScreenshotApp:
    def __init__(self, master):
        self.master = master
        master.title("Employee Monitoring (Local)")
        master.geometry("350x250")

        self.status = "stopped"
        self.thread = None
        
        # --- HR Email Configuration ---
        self.hr_email = "prottoy.saha@soniagroup.com"
        
        # --- !!! SET YOUR PASSWORD HERE !!! ---
        # This password will be required to *edit* the document.
        self.report_password = "MaCroDmT" 
        # ---------------------------------------

        # --- User Name Handling ---
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        self.config_file = os.path.join(self.script_dir, "user_config.txt")
        self.username = None
        
        self.master.withdraw() 
        
        if not self.load_or_request_username():
            return
            
        self.master.deiconify()
        # --------------------------------

        # --- UI Elements ---
        self.greeting_label = tk.Label(master, text="", font=("Arial", 12, "bold"))
        self.greeting_label.pack(pady=(10, 0))

        self.time_label = tk.Label(master, text="", font=("Arial", 10))
        self.time_label.pack(pady=5)
        
        self.time_thread = threading.Thread(target=self.update_time_greeting, daemon=True)
        self.time_thread.start()

        # --- Buttons with Colors ---
        self.start_btn = tk.Button(master, text="Start", width=10, command=self.start, 
                                   bg="#4CAF50", fg="white", font=("Arial", 10, "bold"))
        self.start_btn.pack(pady=5)

        self.pause_btn = tk.Button(master, text="Pause", width=10, command=self.pause, 
                                   bg="#FFEB3B", fg="black", font=("Arial", 10, "bold"))
        self.pause_btn.pack(pady=5)

        self.stop_btn = tk.Button(master, text="Stop", width=10, command=self.stop, 
                                  bg="#F44336", fg="white", font=("Arial", 10, "bold"))
        self.stop_btn.pack(pady=5)

        self.screenshot_dir = os.path.join(self.script_dir, "screenshots")
        if not os.path.exists(self.screenshot_dir):
            os.makedirs(self.screenshot_dir)

    def load_or_request_username(self):
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, "r") as f:
                    self.username = f.read().strip()
                if not self.username:
                    raise FileNotFoundError
            else:
                raise FileNotFoundError
        except FileNotFoundError:
            name = ""
            while not name:
                name = simpledialog.askstring("Welcome!", 
                                              "This is your first time. Please enter your name:", 
                                              parent=self.master)
                if name is None:
                    self.master.destroy()
                    return False
            
            self.username = name
            try:
                with open(self.config_file, "w") as f:
                    f.write(self.username)
            except Exception as e:
                messagebox.showerror("Error", f"Could not save your name: {e}")
        
        return True

    def update_time_greeting(self):
        while True:
            now = datetime.now()
            hour = now.hour
            
            if 5 <= hour < 12:
                greeting = f"Good Morning, {self.username}"
            elif 12 <= hour < 14:
                greeting = f"Good Noon, {self.username}"
            elif 14 <= hour < 18:
                greeting = f"Good Afternoon, {self.username}"
            else:
                greeting = f"Good Evening, {self.username}"
                
            time_str = now.strftime('%Y-%m-%d %H:%M:%S')
            
            try:
                self.greeting_label.config(text=greeting)
                self.time_label.config(text=time_str)
            except tk.TclError:
                break
            
            time.sleep(1)

    def capture_screenshots(self):
        while self.status == "started":
            now = datetime.now()
            filename = f"screenshot_{now.strftime('%Y%m%d_%H%M%S')}.png"
            filepath = os.path.join(self.screenshot_dir, filename)
            
            try:
                screenshot = pyautogui.screenshot()
                screenshot.save(filepath)
                print(f"Saved: {filename}")
            except Exception as e:
                print(f"Error taking screenshot: {e}")

            for _ in range(300): # 5-minute wait
                if self.status != "started":
                    break
                time.sleep(1)

    def start(self):
        if self.status != "started":
            self.status = "started"
            if self.thread is None or not self.thread.is_alive():
                self.thread = threading.Thread(target=self.capture_screenshots)
                self.thread.start()
            messagebox.showinfo("Status", "Screenshot capture started.")

    def pause(self):
        if self.status == "started":
            self.status = "paused"
            messagebox.showinfo("Status", "Screenshot capture paused.")

    def stop(self):
        if self.status != "stopped":
            self.status = "stopped"
            if self.thread and self.thread.is_alive():
                self.thread.join()
            
            # Create and protect the Word document
            file_path = self.create_word_document()
            
            if file_path:
                # If file was created, try to send it
                self.send_email_via_outlook(file_path)
            else:
                # No file was created
                messagebox.showinfo("Status", "Screenshot capture stopped. No screenshots were taken today.")

    def create_word_document(self):
        """
        Creates the Word document, applies edit protection, and cleans up PNGs.
        Returns the path to the DOCX file on success, or None on failure.
        """
        today_str = date.today().strftime("%Y%m%d")
        
        doc_filename = f"screenshots_{self.username}_{today_str}.docx"
        downloads_path = os.path.join(os.path.expanduser("~"), "Downloads", doc_filename)

        today_files = []
        for file in os.listdir(self.screenshot_dir):
            if file.startswith(f"screenshot_{today_str}") and file.endswith(".png"):
                today_files.append(os.path.join(self.screenshot_dir, file))
        
        if not today_files:
            print("No screenshots taken today to save.")
            return None  # Return None if no files

        today_files.sort()
        document = Document()
        print(f"Creating Word document with {len(today_files)} images...")

        # ... (Loop to add images - same as before) ...
        for i in range(0, len(today_files), 2):
            if i > 0:
                document.add_page_break()
            img_path1 = today_files[i]
            try:
                p = document.add_paragraph()
                r = p.add_run()
                r.add_picture(img_path1, width=Inches(6.0)) 
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error adding image {img_path1}: {e}")
                r.add_text(f"[Error: Could not load image {os.path.basename(img_path1)}]")

            if (i + 1) < len(today_files):
                img_path2 = today_files[i + 1]
                try:
                    p = document.add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_path2, width=Inches(6.0))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image {img_path2}: {e}")
                    r.add_text(f"[Error: Could not load image {os.path.basename(img_path2)}]")

        try:
            # 1. Save the document normally from python-docx
            document.save(downloads_path)
            print(f"Word document saved to: {downloads_path}")

            # --- NEW: Apply Read-Only Protection ---
            self.apply_protection(downloads_path)
            # ---------------------------------------
            
            print("Cleaning up temporary PNG files...")
            for img_path in today_files:
                try:
                    os.remove(img_path)
                except Exception as e:
                    print(f"Warning: Could not remove file {img_path}. {e}")
            print("Cleanup complete.")
            
            return downloads_path  # Return the file path on success
            
        except Exception as e:
            print(f"Error saving/protecting Word document: {e}")
            messagebox.showerror("Error", f"Could not save or protect Word file.\n\n{e}")
            return None  # Return None on failure

    # --- NEW: Function to apply protection ---
    def apply_protection(self, file_path):
        """Uses pywin32 to apply edit restrictions to the Word file."""
        word_app = None
        doc = None
        try:
            # Start Word application in the background
            word_app = win32com.client.Dispatch('Word.Application')
            word_app.Visible = False
            
            # Open the document
            doc = word_app.Documents.Open(file_path)
            
            # Apply protection: 
            # Type=2 means wdAllowOnlyReading (allow read-only access)
            doc.Protect(Type=2, NoReset=True, Password=self.report_password)
            
            # Save and close the document
            doc.Save()
            doc.Close()
            
            print(f"Successfully applied read-only protection to {file_path}")

        except Exception as e:
            print(f"Error applying Word protection: {e}")
            messagebox.showerror("Protection Error", 
                                 "Could not apply read-only protection to Word file.\n"
                                 "Please ensure Microsoft Word is installed.\n\n"
                                 f"Error: {e}")
        finally:
            # CRITICAL: Always close the Word application
            if doc:
                doc = None
            if word_app:
                word_app.Quit()
                word_app = None
    # -------------------------------------------

    # --- Function to Send Email (from previous step) ---
    def send_email_via_outlook(self, attachment_path):
        """
        Creates and sends an email with the specified attachment
        using the local Outlook application.
        """
        print(f"Attempting to send email to {self.hr_email}...")
        try:
            outlook = win32com.client.Dispatch('Outlook.Application')
            mail = outlook.CreateItem(0)  # 0 = MailItem
            
            mail.To = self.hr_email
            
            today_str_formatted = date.today().strftime("%Y-%m-%d")
            mail.Subject = f"Screenshot Report - {self.username} - {today_str_formatted}"
            
            mail.Body = (
                f"Hello,\n\n"
                f"Please find the attached screenshot report for {self.username} for {today_str_formatted}.\n\n"
                f"This file is read-only. Editing has been restricted.\n\n"
                f"This is an automated message.\n"
            )
            
            mail.Attachments.Add(attachment_path)
            
            # Use .Send() to send automatically
            mail.Send()
            
            print("Email sent successfully.")
            messagebox.showinfo("Email Status", "Report was successfully created, protected, and sent to the HR Manager.")
            
        except Exception as e:
            print(f"Error sending email: {e}")
            messagebox.showerror("Email Error", 
                                 "Could not send the email. Please ensure Outlook is installed.\n\n"
                                 "If Outlook is open, you may need to 'Allow' the security prompt.\n\n"
                                 f"Error: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = ScreenshotApp(root)
    if app.username: 
        root.mainloop()