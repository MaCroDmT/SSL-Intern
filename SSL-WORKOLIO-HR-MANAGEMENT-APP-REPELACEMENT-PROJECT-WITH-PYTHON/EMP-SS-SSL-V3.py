import tkinter as tk
from tkinter import messagebox, simpledialog
import threading
import pyautogui
import time
import os
from datetime import datetime, date

# --- Imports for Word Document ---
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: 'python-docx' library not found.")
    print("Please install it using: pip install python-docx")
    exit()

# --- Imports for PDF Conversion & Encryption ---
try:
    from docx2pdf import convert
except ImportError:
    print("Error: 'docx2pdf' library not found.")
    print("Please install it using: pip install docx2pdf")
    print("NOTE: docx2pdf also requires MS Word (Windows) or LibreOffice (Mac/Linux) to be installed.")
    exit()

try:
    from pypdf import PdfWriter, PdfReader
except ImportError:
    print("Error: 'pypdf' library not found.")
    print("Please install it using: pip install pypdf")
    exit()

# --- NEW: Import for Outlook Automation ---
try:
    import win32com.client
except ImportError:
    print("Error: 'pywin32' library not found.")
    print("Please install it using: pip install pywin32")
    exit()
# --------------------------------------------------


class ScreenshotApp:
    def __init__(self, master):
        self.master = master
        master.title("Employee Monitoring (Local)")
        master.geometry("350x250")

        self.status = "stopped"
        self.thread = None
        
        # --- Hard-code the encryption password ---
        self.report_password = "MaCroDmT" 
        
        # --- NEW: HR Email Configuration ---
        self.hr_email = "prottoy.saha@soniagroup.com"
        # ------------------------------------

        # --- User Name Handling ---
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        self.config_file = os.path.join(self.script_dir, "user_config.txt")
        self.username = None
        
        self.master.withdraw() 
        
        if not self.load_or_request_username():
            return
            
        self.master.deiconify()
        # --------------------------------

        # --- UI Elements ---
        self.greeting_label = tk.Label(master, text="", font=("Arial", 12, "bold"))
        self.greeting_label.pack(pady=(10, 0))

        self.time_label = tk.Label(master, text="", font=("Arial", 10))
        self.time_label.pack(pady=5)
        
        self.time_thread = threading.Thread(target=self.update_time_greeting, daemon=True)
        self.time_thread.start()

        # --- Buttons with Colors ---
        self.start_btn = tk.Button(master, text="Start", width=10, command=self.start, 
                                   bg="#4CAF50", fg="white", font=("Arial", 10, "bold"))
        self.start_btn.pack(pady=5)

        self.pause_btn = tk.Button(master, text="Pause", width=10, command=self.pause, 
                                   bg="#FFEB3B", fg="black", font=("Arial", 10, "bold"))
        self.pause_btn.pack(pady=5)

        self.stop_btn = tk.Button(master, text="Stop", width=10, command=self.stop, 
                                  bg="#F44336", fg="white", font=("Arial", 10, "bold"))
        self.stop_btn.pack(pady=5)

        # Folder to save screenshots
        self.screenshot_dir = os.path.join(self.script_dir, "screenshots")
        if not os.path.exists(self.screenshot_dir):
            os.makedirs(self.screenshot_dir)

    def load_or_request_username(self):
        """Loads username from config file or prompts for it."""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, "r") as f:
                    self.username = f.read().strip()
                if not self.username:
                    raise FileNotFoundError
            else:
                raise FileNotFoundError
        except FileNotFoundError:
            name = ""
            while not name:
                name = simpledialog.askstring("Welcome!", 
                                              "This is your first time. Please enter your name:", 
                                              parent=self.master)
                if name is None:
                    self.master.destroy()
                    return False
            
            self.username = name
            try:
                with open(self.config_file, "w") as f:
                    f.write(self.username)
            except Exception as e:
                messagebox.showerror("Error", f"Could not save your name: {e}")
        
        return True

    def update_time_greeting(self):
        """Continuously updates the time and greeting message."""
        while True:
            now = datetime.now()
            hour = now.hour
            
            if 5 <= hour < 12:
                greeting = f"Good Morning, {self.username}"
            elif 12 <= hour < 14:
                greeting = f"Good Noon, {self.username}"
            elif 14 <= hour < 18:
                greeting = f"Good Afternoon, {self.username}"
            else:
                greeting = f"Good Evening, {self.username}"
                
            time_str = now.strftime('%Y-%m-%d %H:%M:%S')
            
            try:
                self.greeting_label.config(text=greeting)
                self.time_label.config(text=time_str)
            except tk.TclError:
                break
            
            time.sleep(1)

    def capture_screenshots(self):
        """The main loop for capturing screenshots."""
        while self.status == "started":
            now = datetime.now()
            filename = f"screenshot_{now.strftime('%Y%m%d_%H%M%S')}.png"
            filepath = os.path.join(self.screenshot_dir, filename)
            
            try:
                screenshot = pyautogui.screenshot()
                screenshot.save(filepath)
                print(f"Saved: {filename}")
            except Exception as e:
                print(f"Error taking screenshot: {e}")

            for _ in range(300): # 5-minute wait
                if self.status != "started":
                    break
                time.sleep(1)

    def start(self):
        if self.status != "started":
            self.status = "started"
            if self.thread is None or not self.thread.is_alive():
                self.thread = threading.Thread(target=self.capture_screenshots)
                self.thread.start()
            messagebox.showinfo("Status", "Screenshot capture started.")

    def pause(self):
        if self.status == "started":
            self.status = "paused"
            messagebox.showinfo("Status", "Screenshot capture paused.")

    # --- MODIFIED: Stop function now creates report AND emails it ---
    def stop(self):
        if self.status != "stopped":
            self.status = "stopped"
            if self.thread and self.thread.is_alive():
                self.thread.join()
            
            # Pass the hard-coded password to the report creator
            success, final_path = self.create_report(self.report_password)
            
            if success and final_path != "No screenshots to save":
                # --- NEW: Call the email function on success ---
                print(f"Report created. Now emailing {final_path}...")
                self.send_email_via_outlook(final_path)
                # The email function will show its own success/error message
                # ------------------------------------------------
            elif final_path == "No screenshots to save":
                 messagebox.showinfo("Status", "Capture stopped. No new screenshots to save.")
            else:
                 # Error message is already shown by create_report
                 print("Failed to create report.")


    def cleanup_pngs(self, png_file_list):
        """Deletes the list of PNG files provided."""
        print("Cleaning up temporary PNG files...")
        for img_path in png_file_list:
            try:
                os.remove(img_path)
            except Exception as e:
                print(f"Warning: Could not remove file {img_path}. {e}")
        print("Cleanup complete.")
        

    def create_report(self, password):
        """Finds today's screenshots and saves them to an encrypted PDF.
           Returns (True, file_path) on success, (False, None) on failure.
        """
        today_str = date.today().strftime("%Y%m%d")
        
        # --- Find today's files ---
        today_files = []
        for file in os.listdir(self.screenshot_dir):
            if file.startswith(f"screenshot_{today_str}") and file.endswith(".png"):
                today_files.append(os.path.join(self.screenshot_dir, file))
        
        if not today_files:
            print("No screenshots taken today to save.")
            return (True, "No screenshots to save")

        today_files.sort()
        document = Document()
        print(f"Creating report with {len(today_files)} images...")

        # --- Build the document content ---
        for i in range(0, len(today_files), 2):
            if i > 0:
                document.add_page_break()

            img_path1 = today_files[i]
            try:
                p = document.add_paragraph()
                r = p.add_run()
                r.add_picture(img_path1, width=Inches(6.0)) 
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error adding image {img_path1}: {e}")
                r.add_text(f"[Error: Could not load image {os.path.basename(img_path1)}]")

            if (i + 1) < len(today_files):
                img_path2 = today_files[i + 1]
                try:
                    p = document.add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_path2, width=Inches(6.0))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image {img_path2}: {e}")
                    r.add_text(f"[Error: Could not load image {os.path.basename(img_path2)}]")

        # --- Define file paths ---
        base_filename = f"screenshots_{self.username}_{today_str}"
        temp_doc_path = os.path.join(self.screenshot_dir, f"{base_filename}_temp.docx")
        temp_pdf_path = os.path.join(self.screenshot_dir, f"{base_filename}_temp.pdf")
        final_pdf_path = os.path.join(os.path.expanduser("~"), "Downloads", f"{base_filename}_SECURE.pdf")

        try:
            # 1. Save the DOCX temporarily
            document.save(temp_doc_path)

            # 2. Convert DOCX to PDF
            print(f"Converting {temp_doc_path} to PDF...")
            convert(temp_doc_path, temp_pdf_path)

            # 3. Encrypt the PDF
            print(f"Encrypting PDF to {final_pdf_path}...")
            reader = PdfReader(temp_pdf_path)
            writer = PdfWriter()

            for page in reader.pages:
                writer.add_page(page)

            writer.encrypt(password) 

            with open(final_pdf_path, "wb") as f:
                writer.write(f)
            
            # 4. Clean up temp DOCX and temp PDF
            os.remove(temp_doc_path)
            os.remove(temp_pdf_path)
            
            # 5. Clean up PNGs
            self.cleanup_pngs(today_files)
            return (True, final_pdf_path)
                 
        except Exception as e:
            print(f"Error saving or converting document: {e}")
            messagebox.showerror("Error", f"Could not save report file.\nIs it open?\n\n{e}")
            # Clean up temp files if they exist
            if os.path.exists(temp_doc_path):
                try: os.remove(temp_doc_path); 
                except: pass
            if os.path.exists(temp_pdf_path):
                try: os.remove(temp_pdf_path);
                except: pass
            return (False, None)

    # --- NEW: Function to Send Email via Outlook ---
    def send_email_via_outlook(self, attachment_path):
        """
        Creates and sends an email with the specified PDF attachment
        using the local Outlook application.
        """
        print(f"Attempting to send email to {self.hr_email}...")
        try:
            outlook = win32com.client.Dispatch('Outlook.Application')
            mail = outlook.CreateItem(0)  # 0 = MailItem
            
            mail.To = self.hr_email
            
            today_str_formatted = date.today().strftime("%Y-%m-%d")
            mail.Subject = f"Screenshot Report - {self.username} - {today_str_formatted}"
            
            # Updated body to mention the encrypted PDF
            mail.Body = (
                f"Hello,\n\n"
                f"Please find the attached screenshot report for {self.username} for {today_str_formatted}.\n\n"
                f"**This is an encrypted PDF file.**\n\n"
                # f"(Password: {self.report_password})\n\n"
                f"This is an automated message.\n"
            )
            
            mail.Attachments.Add(attachment_path)
            
            # Use .Send() to send automatically
            mail.Send()
            
            print("Email sent successfully.")
            messagebox.showinfo("Email Status", "Report was successfully created, encrypted, and sent to the HR Manager.")
            
        except Exception as e:
            print(f"Error sending email: {e}")
            messagebox.showerror("Email Error", 
                                 "Could not send the email. Please ensure Outlook is installed.\n\n"
                                 "If Outlook is open, you may need to 'Allow' the security prompt.\n\n"
                                 f"Error: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = ScreenshotApp(root)
    if app.username: 
        root.mainloop()