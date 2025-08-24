import xlwings as xw
import pandas as pd
from docx import Document
from tkinter import Tk, filedialog
import os

# --- Select Excel file ---
Tk().withdraw()
excel_file = filedialog.askopenfilename(
    title="Select Excel File",
    filetypes=[("Excel files", "*.xlsx *.xlsm *.xlsb *.xls")]
)
if not excel_file:
    print("❌ No file selected.")
    exit()

# --- Open workbook ---
app = xw.App(visible=False)
wb = app.books.open(excel_file)

doc = Document()

# Loop sheets
for sheet in wb.sheets:
    print(f"📑 Processing Sheet: {sheet.name}")
    used_range = sheet.used_range
    values = used_range.value

    if not values:
        continue

    headers = values[0]
    print("Headers:", headers)

    # Find required columns
    try:
        style_photo_col = headers.index("Style Photo")
        style_no_col = headers.index("Style No.")
    except ValueError:
        print("❌ Missing 'Style Photo' or 'Style No.' columns")
        continue

    # Loop rows (skip header)
    for r in range(2, used_range.rows.count + 1):
        for c in range(1, used_range.columns.count + 1):
            cell = sheet.range((r, c))
            color = cell.api.Interior.Color  # Excel RGB color
            
            # Convert color to int if it's a float
            if isinstance(color, float):
                color = int(color)

            # Check if color is a valid integer before processing
            if isinstance(color, int) and color >= 0:
                # Convert to hex (#RRGGBB)
                rgb = f"#{color & 255:02X}{(color >> 8) & 255:02X}{(color >> 16) & 255:02X}"

                if rgb.upper() == "#00FF00":  # Only green highlights
                    style_no = sheet.range((r, style_no_col + 1)).value
                    style_photo = sheet.range((r, style_photo_col + 1)).value
                    col_name = headers[c - 1]
                    highlight_date = cell.value

                    print(f"✅ Green cell at {cell.address}: {rgb}, Style No: {style_no}, Date: {highlight_date}")

                    # --- Build mail template ---
                    doc.add_paragraph(f"Subject: Style No. {style_no}, {col_name}, Date: {highlight_date}")
                    doc.add_paragraph("")
                    doc.add_paragraph("Hi,")
                    doc.add_paragraph(
                        f"This is to let you know style no.: {style_no} has a deadline in {highlight_date} "
                        f"regarding this {col_name}."
                    )
                    doc.add_paragraph("Please, update us as soon as possible.")
                    if style_photo:
                        doc.add_paragraph(f"[Image placeholder for {style_photo}]")
                    doc.add_paragraph("-" * 40)

# --- Save Word file ---
output_file = os.path.join(os.path.dirname(excel_file), "Generated_Mails.docx")
doc.save(output_file)
print(f"📂 Word file generated: {output_file}")

wb.close()
app.quit()
