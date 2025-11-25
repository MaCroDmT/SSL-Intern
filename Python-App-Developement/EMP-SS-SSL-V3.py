import tkinter as tk
from tkinter import messagebox, simpledialog
import threading
import pyautogui
import time
import os
from datetime import datetime, date

# --- Imports for Word Document ---
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: 'python-docx' library not found.")
    print("Please install it using: pip install python-docx")
    exit()

# --- Imports for PDF Conversion & Encryption ---
try:
    from docx2pdf import convert
except ImportError:
    print("Error: 'docx2pdf' library not found.")
    print("Please install it using: pip install docx2pdf")
    print("NOTE: docx2pdf also requires MS Word (Windows) or LibreOffice (Mac/Linux) to be installed.")
    exit()

try:
    from pypdf import PdfWriter, PdfReader
except ImportError:
    print("Error: 'pypdf' library not found.")
    print("Please install it using: pip install pypdf")
    exit()
# --------------------------------------------------


class ScreenshotApp:
    def __init__(self, master):
        self.master = master
        master.title("Employee Monitoring (Local)")
        master.geometry("350x250")

        self.status = "stopped"
        self.thread = None
        
        # --- NEW: Hard-code the encryption password ---
        # !!! IMPORTANT: Change this to your own secret password !!!
        self.report_password = "MaCroDmT" 
        # ------------------------------------------------

        # --- User Name Handling ---
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        self.config_file = os.path.join(self.script_dir, "user_config.txt")
        self.username = None
        
        self.master.withdraw() 
        
        if not self.load_or_request_username():
            return
            
        self.master.deiconify()
        # --------------------------------

        # --- UI Elements ---
        self.greeting_label = tk.Label(master, text="", font=("Arial", 12, "bold"))
        self.greeting_label.pack(pady=(10, 0))

        self.time_label = tk.Label(master, text="", font=("Arial", 10))
        self.time_label.pack(pady=5)
        
        self.time_thread = threading.Thread(target=self.update_time_greeting, daemon=True)
        self.time_thread.start()

        # --- Buttons with Colors ---
        self.start_btn = tk.Button(master, text="Start", width=10, command=self.start, 
                                   bg="#4CAF50", fg="white", font=("Arial", 10, "bold"))
        self.start_btn.pack(pady=5)

        self.pause_btn = tk.Button(master, text="Pause", width=10, command=self.pause, 
                                   bg="#FFEB3B", fg="black", font=("Arial", 10, "bold"))
        self.pause_btn.pack(pady=5)

        self.stop_btn = tk.Button(master, text="Stop", width=10, command=self.stop, 
                                  bg="#F44336", fg="white", font=("Arial", 10, "bold"))
        self.stop_btn.pack(pady=5)

        # Folder to save screenshots
        self.screenshot_dir = os.path.join(self.script_dir, "screenshots")
        if not os.path.exists(self.screenshot_dir):
            os.makedirs(self.screenshot_dir)

    def load_or_request_username(self):
        """
        Loads username from config file. If file doesn't exist,
        prompts user for their name and saves it.
        Returns True on success, False if user cancels.
        """
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, "r") as f:
                    self.username = f.read().strip()
                if not self.username:
                    raise FileNotFoundError
            else:
                raise FileNotFoundError
        except FileNotFoundError:
            name = ""
            while not name:
                name = simpledialog.askstring("Welcome!", 
                                              "This is your first time. Please enter your name:", 
                                              parent=self.master)
                if name is None:
                    self.master.destroy()
                    return False
            
            self.username = name
            try:
                with open(self.config_file, "w") as f:
                    f.write(self.username)
            except Exception as e:
                messagebox.showerror("Error", f"Could not save your name: {e}")
        
        return True

    def update_time_greeting(self):
        """Continuously updates the time and greeting message."""
        while True:
            now = datetime.now()
            hour = now.hour
            
            if 5 <= hour < 12:
                greeting = f"Good Morning, {self.username}"
            elif 12 <= hour < 14:
                greeting = f"Good Noon, {self.username}"
            elif 14 <= hour < 18:
                greeting = f"Good Afternoon, {self.username}"
            else:
                greeting = f"Good Evening, {self.username}"
                
            time_str = now.strftime('%Y-%m-%d %H:%M:%S')
            
            try:
                self.greeting_label.config(text=greeting)
                self.time_label.config(text=time_str)
            except tk.TclError:
                break
            
            time.sleep(1)

    def capture_screenshots(self):
        """The main loop for capturing screenshots."""
        while self.status == "started":
            now = datetime.now()
            filename = f"screenshot_{now.strftime('%Y%m%d_%H%M%S')}.png"
            filepath = os.path.join(self.screenshot_dir, filename)
            
            try:
                screenshot = pyautogui.screenshot()
                screenshot.save(filepath)
                print(f"Saved: {filename}")
            except Exception as e:
                print(f"Error taking screenshot: {e}")

            for _ in range(300): # 5-minute wait
                if self.status != "started":
                    break
                time.sleep(1)

    def start(self):
        if self.status != "started":
            self.status = "started"
            if self.thread is None or not self.thread.is_alive():
                self.thread = threading.Thread(target=self.capture_screenshots)
                self.thread.start()
            messagebox.showinfo("Status", "Screenshot capture started.")

    def pause(self):
        if self.status == "started":
            self.status = "paused"
            messagebox.showinfo("Status", "Screenshot capture paused.")

    # --- MODIFIED: Stop function now uses the hard-coded password ---
    def stop(self):
        if self.status != "stopped":
            
            # REMOVED: The simpledialog.askstring() call is gone.
            
            self.status = "stopped"
            if self.thread and self.thread.is_alive():
                self.thread.join()
            
            # MODIFIED: Pass the hard-coded password from self.report_password
            success, final_path = self.create_report(self.report_password)
            
            if success and final_path != "No screenshots to save":
                messagebox.showinfo("Status", f"Capture stopped and encrypted PDF created:\n{final_path}")
            elif final_path == "No screenshots to save":
                 messagebox.showinfo("Status", "Capture stopped. No new screenshots to save.")
            else:
                # Error message is already shown by create_report
                print("Failed to create report.")


    # --- NEW: Helper function to clean up PNGs ---
    def cleanup_pngs(self, png_file_list):
        """Deletes the list of PNG files provided."""
        print("Cleaning up temporary PNG files...")
        for img_path in png_file_list:
            try:
                os.remove(img_path)
            except Exception as e:
                print(f"Warning: Could not remove file {img_path}. {e}")
        print("Cleanup complete.")
        

    # --- MODIFIED: Renamed and updated to ONLY create encrypted PDF ---
    def create_report(self, password):
        """Finds today's screenshots and saves them to an encrypted PDF.
           Returns (True, file_path) on success, (False, None) on failure.
        """
        today_str = date.today().strftime("%Y%m%d")
        
        # --- Find today's files ---
        today_files = []
        for file in os.listdir(self.screenshot_dir):
            if file.startswith(f"screenshot_{today_str}") and file.endswith(".png"):
                today_files.append(os.path.join(self.screenshot_dir, file))
        
        if not today_files:
            print("No screenshots taken today to save.")
            return (True, "No screenshots to save") # Not a failure, just nothing to do

        today_files.sort()
        document = Document()
        print(f"Creating report with {len(today_files)} images...")

        # --- Build the document content ---
        for i in range(0, len(today_files), 2):
            if i > 0:
                document.add_page_break()

            # Add the first image of the pair
            img_path1 = today_files[i]
            try:
                p = document.add_paragraph()
                r = p.add_run()
                r.add_picture(img_path1, width=Inches(6.0)) 
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error adding image {img_path1}: {e}")
                r.add_text(f"[Error: Could not load image {os.path.basename(img_path1)}]")

            # Check for and add the second image of the pair
            if (i + 1) < len(today_files):
                img_path2 = today_files[i + 1]
                try:
                    p = document.add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_path2, width=Inches(6.0))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image {img_path2}: {e}")
                    r.add_text(f"[Error: Could not load image {os.path.basename(img_path2)}]")

        # --- Define file paths ---
        base_filename = f"screenshots_{self.username}_{today_str}"
        # We'll save the docx temporarily in the app's screenshot dir
        temp_doc_path = os.path.join(self.screenshot_dir, f"{base_filename}_temp.docx")
        temp_pdf_path = os.path.join(self.screenshot_dir, f"{base_filename}_temp.pdf")
        final_pdf_path = os.path.join(os.path.expanduser("~"), "Downloads", f"{base_filename}_SECURE.pdf")

        try:
            # --- This is now the ONLY path ---
            
            # 1. Save the DOCX temporarily
            document.save(temp_doc_path)

            # 2. Convert DOCX to PDF
            print(f"Converting {temp_doc_path} to PDF...")
            convert(temp_doc_path, temp_pdf_path)

            # 3. Encrypt the PDF
            print(f"Encrypting PDF to {final_pdf_path}...")
            reader = PdfReader(temp_pdf_path)
            writer = PdfWriter()

            for page in reader.pages:
                writer.add_page(page)

            # Apply encryption using the provided password
            writer.encrypt(password) 

            with open(final_pdf_path, "wb") as f:
                writer.write(f)
            
            # 4. Clean up temp DOCX and temp PDF
            os.remove(temp_doc_path)
            os.remove(temp_pdf_path)
            
            # 5. Clean up PNGs
            self.cleanup_pngs(today_files)
            return (True, final_pdf_path)
                
        except Exception as e:
            print(f"Error saving or converting document: {e}")
            messagebox.showerror("Error", f"Could not save report file.\nIs it open?\n\n{e}")
            # Clean up temp files if they exist
            if os.path.exists(temp_doc_path):
                try: os.remove(temp_doc_path); 
                except: pass
            if os.path.exists(temp_pdf_path):
                try: os.remove(temp_pdf_path);
                except: pass
            return (False, None)


if __name__ == "__main__":
    root = tk.Tk()
    app = ScreenshotApp(root)
    if app.username: 
        root.mainloop()