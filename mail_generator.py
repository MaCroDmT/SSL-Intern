import pandas as pd
import openpyxl
import os
from docx import Document
from docx.shared import Inches
from tkinter import Tk, filedialog, messagebox

# =======================
# Step 1: Select Excel File via Tkinter
# =======================
root = Tk()
root.withdraw()  # Hide main window
messagebox.showinfo("Excel Mail Generator", "Please select your Excel file.")
file_path = filedialog.askopenfilename(
    title="Select Excel File",
    filetypes=[("Excel files", "*.xlsx *.xls")]
)

if not file_path:
    messagebox.showerror("Error", "No file selected.")
    exit()

print("✅ Selected file:", file_path)

# =======================
# Step 2: Extract Images from Excel
# =======================
wb = openpyxl.load_workbook(file_path)
ws = wb.active

os.makedirs("excel_images", exist_ok=True)

images = []
for idx, image in enumerate(ws._images, start=1):
    img_bytes = image._data()
    img_path = f"excel_images/image_{idx}.png"
    with open(img_path, "wb") as f:
        f.write(img_bytes)

    # openpyxl: row is 0-based, Excel is 1-based, DF index = ExcelRow - 2
    excel_row = image.anchor._from.row + 1
    df_row = excel_row - 2
    images.append({
        "col": image.anchor._from.col,
        "excel_row": excel_row,
        "df_row": df_row,
        "path": img_path
    })

print("✅ Extracted images:", [(img["col"], img["excel_row"], img["path"]) for img in images])

# =======================
# Step 3: Load Excel into DataFrame
# =======================
df = pd.read_excel(file_path)
print("✅ Data loaded:")
print(df.head())

# =======================
# Step 4: Find '*' Cells
# =======================
asterisk_cells = []
for col in df.columns:
    for row_idx, cell in enumerate(df[col]):
        if isinstance(cell, str) and '*' in cell:
            asterisk_cells.append({
                "row_idx": row_idx,
                "col_name": col,
                "cell_value": cell
            })

print(f"✅ Found {len(asterisk_cells)} asterisk-marked cells")

# =======================
# Step 5: Find Style Photo Column
# =======================
style_photo_col_idx = None
for i, col in enumerate(df.columns):
    if "Style" in col and "Photo" in col:
        style_photo_col_idx = i
        break

# =======================
# Step 6: Generate Mail Templates DOCX
# =======================
doc = Document()
doc.add_heading("Generated Mail Templates (Only Asterisk Marked Cells)", 0)

for cell_info in asterisk_cells:
    row_idx = cell_info["row_idx"]
    col_name = cell_info["col_name"]
    cell_value = cell_info["cell_value"]

    style_no = str(df.loc[row_idx, "Style No."]) if "Style No." in df.columns else "Unknown"
    if pd.isna(style_no) or style_no.strip() == "":
        style_no = "Unknown"

    style_photo_path = None
    if style_photo_col_idx is not None:
        candidates = [img for img in images if img["col"] == style_photo_col_idx]
        if candidates:
            nearest_img = min(candidates, key=lambda img: abs(img["df_row"] - row_idx))
            style_photo_path = nearest_img["path"]

    # Construct mail text
    subject = f"Regarding Style No: {style_no}, {col_name}, Date: {cell_value.replace('*','')}"
    body = f"""
Hi,

This is to let you know regarding style no.: {style_no} has a deadline in: {cell_value.replace('*','')} regarding this: {col_name}.

Please, update us as soon as possible.

Thank You
"""
    doc.add_heading(subject, level=1)
    doc.add_paragraph(body)
    if style_photo_path:
        doc.add_picture(style_photo_path, width=Inches(2))
    else:
        doc.add_paragraph("[⚠️ No image found for this row]")
    doc.add_paragraph("\n" + "-"*50 + "\n")

# =======================
# Step 7: Save DOCX & Notify
# =======================
output_path = os.path.join(os.path.dirname(file_path), "mail_templates_starred.docx")
doc.save(output_path)

print(f"✅ Mail templates saved to {output_path}")
messagebox.showinfo("Done", f"Mail templates saved to:\n{output_path}")
