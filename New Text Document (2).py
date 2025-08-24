import xlwings as xw
from docx import Document
from docx.shared import Inches
import os
from tkinter import Tk, filedialog

# === File Picker ===
Tk().withdraw()  # Hide root window
EXCEL_FILE = filedialog.askopenfilename(
    title="Select Excel File",
    filetypes=[("Excel Files", "*.xlsx *.xlsm *.xlsb *.xls")]
)

if not EXCEL_FILE:
    print("❌ No file selected, exiting...")
    exit()

OUTPUT_DOCX = "Generated_Mails.docx"
GREEN_HEX = "00FF00"  # Green highlight color

# === Open Excel ===
app = xw.App(visible=False)
wb = app.books.open(EXCEL_FILE)
ws = wb.sheets[0]  # First sheet, adjust if needed

doc = Document()

# === Read headers ===
headers = ws.range("A1").expand("right").value
print("Headers:", headers)

# Find column indexes
style_photo_idx = headers.index("Style Photo") + 1
style_no_idx = headers.index("Style No.") + 1

# === Iterate cells ===
used_range = ws.used_range
for cell in used_range:
    color = cell.color  # RGB tuple or None
    if color:
        r, g, b = color
        hex_code = f"{r:02X}{g:02X}{b:02X}"

        if hex_code == GREEN_HEX:  # Match highlight
            row = cell.row
            col = cell.column

            column_name = headers[col - 1]
            deadline_date = cell.value
            style_no = ws.cells(row, style_no_idx).value
            style_photo = ws.cells(row, style_photo_idx).value  # Must be file path

            # --- Mail Template ---
            subject = f"Style No. {style_no}, column name: {column_name}, Date: {deadline_date}"
            body = f"""
Hi,

This is to let you know style no. : {style_no} has a deadline in {deadline_date} regarding {column_name}.

Please, update us as soon as possible.
"""

            doc.add_heading(subject, level=1)
            doc.add_paragraph(body)

            # Add photo if valid path
            if style_photo and os.path.exists(style_photo):
                try:
                    doc.add_picture(style_photo, width=Inches(2))
                except Exception as e:
                    doc.add_paragraph(f"[Image error: {e}]")

            doc.add_page_break()

# === Save Word ===
doc.save(OUTPUT_DOCX)
wb.close()
app.quit()

print(f"✅ Word file generated: {OUTPUT_DOCX}")
